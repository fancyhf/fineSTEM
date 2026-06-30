"""
文档生成服务

用途：生成项目文档（MD/JSON/PDF/DOCX）与下载元信息
维护者：AI Agent
"""

from __future__ import annotations

import io
import json
import zipfile
from pathlib import Path
from typing import Any

from app.core.config import settings
from app.core.time_utils import utc_now_iso, utc_now
from app.repositories.runtime_db import db
from app.schemas.documents import DocumentFormat, DocumentType


class DocumentService:
    def __init__(self) -> None:
        self.base_path = Path(settings.STORAGE_BASE_PATH)
        self.exports_dir = self.base_path / settings.STORAGE_EXPORTS_DIR
        self.exports_dir.mkdir(parents=True, exist_ok=True)

    def _snapshot(self, project_id: str) -> dict[str, Any]:
        project = db.get_project(project_id)
        if not project:
            raise ValueError("项目不存在")
        progress = db.get_skill_state(project_id)
        evidence_list = db.list_evidence_by_project(project_id, skip=0, limit=1000)
        achievement = db.get_achievement_card_by_project(project_id)
        return {
            "project": project.model_dump(mode="json"),
            "progress": progress.model_dump(mode="json") if progress else {},
            "evidence": [item.model_dump(mode="json") for item in evidence_list],
            "achievement_card": achievement.model_dump(mode="json") if achievement else None,
            "generated_at": utc_now_iso(),
        }

    def _build_markdown(self, document_type: DocumentType, snapshot: dict[str, Any]) -> str:
        project = snapshot["project"]
        progress = snapshot["progress"]
        evidence = snapshot["evidence"]
        achievement = snapshot["achievement_card"]

        title_map: dict[DocumentType, str] = {
            "proposal": "开题报告",
            "technical": "技术报告",
            "final": "结题报告",
        }
        title = f"{project['name']} - {title_map[document_type]}"
        lines = [
            f"# {title}",
            "",
            "## 项目概况",
            f"- 项目ID: {project.get('id', '')}",
            f"- 模式: {project.get('mode', '')}",
            f"- 当前阶段: {project.get('current_stage', '')}",
            f"- 生成时间: {snapshot.get('generated_at', '')}",
            "",
            "## 项目描述",
            project.get("description", "") or "暂无",
            "",
            "## 进度摘要",
            f"- 当前阶段: {progress.get('current_stage', '未知')}",
            f"- 历史节点数: {len(progress.get('stage_history', []))}",
            "",
            "## 证据清单",
        ]
        if evidence:
            for item in evidence[:50]:
                lines.append(
                    f"- [{item.get('type', 'unknown')}] {item.get('content', '')[:120]}"
                )
        else:
            lines.append("- 暂无证据")

        if achievement:
            lines.extend(
                [
                    "",
                    "## 成果档案卡",
                    f"- 标题: {achievement.get('title', '')}",
                    f"- 一句话介绍: {achievement.get('one_liner', '')}",
                    f"- 反思: {achievement.get('reflection', '')}",
                ]
            )

        if document_type == "proposal":
            lines.extend(
                [
                    "",
                    "## 开题重点",
                    "- 研究问题与目标",
                    "- 成功标准与风险",
                    "- 阶段计划与资源评估",
                ]
            )
        elif document_type == "technical":
            lines.extend(
                [
                    "",
                    "## 技术重点",
                    "- 技术路线与架构",
                    "- 关键实现与调试记录",
                    "- 验证方式与结果",
                ]
            )
        else:
            lines.extend(
                [
                    "",
                    "## 结题重点",
                    "- 成果与影响",
                    "- 学习反思",
                    "- 后续迭代建议",
                ]
            )
        return "\n".join(lines) + "\n"

    def _to_docx(self, markdown_text: str) -> bytes:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(4)

        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        for raw in markdown_text.splitlines():
            line = raw.strip()
            if not line:
                continue
            if line.startswith("# "):
                text = line[2:]
                p = doc.add_heading(text, level=1)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            elif line.startswith("## "):
                text = line[3:]
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif line.startswith("- "):
                text = line[2:]
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _to_pdf(self, markdown_text: str) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        font_paths = [
            Path("C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simsun.ttc"),
            Path("C:/Windows/Fonts/simhei.ttf"),
        ]
        cn_font_name = "Helvetica"
        for fp in font_paths:
            if fp.exists():
                try:
                    pdfmetrics.registerFont(TTFont("CNFont", str(fp)))
                    cn_font_name = "CNFont"
                    break
                except Exception:
                    continue

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        style_h1 = ParagraphStyle(
            "CN_H1", parent=styles["Heading1"],
            fontName=cn_font_name, fontSize=18, leading=24,
            spaceAfter=12, alignment=TA_CENTER,
        )
        style_h2 = ParagraphStyle(
            "CN_H2", parent=styles["Heading2"],
            fontName=cn_font_name, fontSize=14, leading=20,
            spaceAfter=8, spaceBefore=12,
        )
        style_normal = ParagraphStyle(
            "CN_Normal", parent=styles["Normal"],
            fontName=cn_font_name, fontSize=11, leading=18,
            spaceAfter=4,
        )
        style_bullet = ParagraphStyle(
            "CN_Bullet", parent=styles["Normal"],
            fontName=cn_font_name, fontSize=11, leading=18,
            leftIndent=20, spaceAfter=2,
        )

        story: list[Any] = []
        for raw in markdown_text.splitlines():
            line = raw.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            if line.startswith("# "):
                text = line[2:]
                story.append(Paragraph(text, style_h1))
            elif line.startswith("## "):
                text = line[3:]
                story.append(Paragraph(text, style_h2))
            elif line.startswith("- "):
                text = line[2:]
                safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(f"\u2022 {safe}", style_bullet))
            else:
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, style_normal))

        doc.build(story)
        return buffer.getvalue()

    def generate(
        self, project_id: str, document_type: DocumentType, fmt: DocumentFormat
    ) -> tuple[str, str, bytes | str]:
        snapshot = self._snapshot(project_id)
        markdown_text = self._build_markdown(document_type, snapshot)
        ts = utc_now().strftime("%Y%m%d%H%M%S")
        basename = f"{project_id}_{document_type}_{ts}"

        if fmt == "md":
            return f"{basename}.md", "text/markdown; charset=utf-8", markdown_text
        if fmt == "json":
            payload = {
                "document_type": document_type,
                "snapshot": snapshot,
                "markdown": markdown_text,
            }
            return (
                f"{basename}.json",
                "application/json; charset=utf-8",
                json.dumps(payload, ensure_ascii=False, indent=2),
            )
        if fmt == "docx":
            return (
                f"{basename}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                self._to_docx(markdown_text),
            )
        if fmt == "pdf":
            return f"{basename}.pdf", "application/pdf", self._to_pdf(markdown_text)
        raise ValueError("不支持的导出格式")


document_service = DocumentService()
